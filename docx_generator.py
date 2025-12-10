from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import base64

def create_table_from_lines(doc, table_lines):
    """Create a properly formatted table from pipe-separated lines"""
    if not table_lines:
        return
        
    # Parse the first line to determine column count
    first_line_cells = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
    if len(first_line_cells) < 2:
        # Not a valid table, add as regular text
        for line in table_lines:
            doc.add_paragraph(line)
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=len(first_line_cells))
    table.style = 'Table Grid'
    
    # Set headers from first line
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(first_line_cells):
        if idx < len(hdr_cells):
            hdr_cells[idx].text = header
            # Make header bold
            for paragraph in hdr_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
    
    # Add data rows (skip first line as it's the header)
    for line in table_lines[1:]:
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if len(cells) == len(first_line_cells):
            row_cells = table.add_row().cells
            for idx, value in enumerate(cells):
                if idx < len(row_cells):
                    row_cells[idx].text = value

def generate_docx_from_translation(translated_text, source_lang="English", target_lang="Hindi"):
    """Generate professional DOCX document from translation results"""
    
    doc = Document()
    
    # Set document margins for professional appearance
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # HEADER WITH COMPANY BRANDING
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = header_para.add_run()
    run.add_text("PNB Housing Finance Ltd.")
    run.font.size = Pt(12)
    run.font.bold = True
    
    # Extract document title from translated text
    lines = [line.strip() for line in translated_text.split('\n') if line.strip()]
    doc_title = "Professional Translation Document"
    
    # Look for title in first few lines
    for line in lines[:3]:
        if line and (line.startswith('#') or len(line) < 80):
            doc_title = line.replace('#', '').strip()
            break
    
    # DOCUMENT TITLE
    title_para = doc.add_heading(doc_title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.size = Pt(18)
    
    # TRANSLATION METADATA SECTION
    doc.add_heading("Translation Information", level=2)
    
    # Create professional metadata table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = 'Table Grid'
    
    # Metadata content
    metadata = [
        ("Source Language:", source_lang),
        ("Target Language:", target_lang),
        ("Translation Method:", "AI-Powered Professional Translation"),
        ("Processing Engine:", "AWS Bedrock Claude 4.5"),
        ("Document Status:", "Verified Translation")
    ]
    
    for i, (label, value) in enumerate(metadata):
        meta_table.cell(i, 0).text = label
        meta_table.cell(i, 1).text = value
        
        # Make labels bold
        for paragraph in meta_table.cell(i, 0).paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add spacing
    doc.add_paragraph()
    
    # TRANSLATED CONTENT SECTION
    doc.add_heading("Translated Document", level=2)
    
    # Process the translated text with proper formatting
    i = 0
    while i < len(lines):
        line = lines[i]
        
        if not line:
            doc.add_paragraph()
            i += 1
            continue
            
        # Handle headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        # Handle numbered lists
        elif line.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            p = doc.add_paragraph(line.strip())
            p.style = 'List Number'
        # Handle bullet points
        elif line.strip().startswith(('•', '-', '*')):
            p = doc.add_paragraph(line.strip()[1:].strip())
            p.style = 'List Bullet'
        # Handle tables (pipe-separated format)
        elif '|' in line and len(line.split('|')) > 2:
            # Collect all consecutive table lines
            table_lines = []
            j = i
            while j < len(lines) and lines[j] and '|' in lines[j] and len(lines[j].split('|')) > 2:
                table_lines.append(lines[j])
                j += 1
            
            if table_lines:
                # Process the table
                create_table_from_lines(doc, table_lines)
                i = j - 1  # Skip processed lines
        # Regular paragraphs
        else:
            p = doc.add_paragraph()
            
            # Handle bold text
            parts = line.split('**')
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:  # Odd indices are bold
                    run.font.bold = True
            
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        i += 1
    
    # FOOTER WITH CONFIDENTIALITY NOTICE
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "CONFIDENTIAL DOCUMENT | PNB Housing Finance Ltd. | Professional Translation Service"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    
    # Return as bytes for web response
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()

# Legacy function for backward compatibility
def generate_docx(title="Professional Translation Document", 
                 body="", 
                 logo_path=None,
                 output_path="output.docx"):
    """Generate a professional DOCX document with header, body, table, and footer"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # HEADER WITH LOGO
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = header_para.add_run()
    if logo_path:
        try:
            run.add_picture(logo_path, width=Inches(1.2))
        except:
            run.add_text("[Logo missing]")
    else:
        # Add company name as text logo
        run.add_text("PNB Housing Finance Ltd.")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = None
    
    # TITLE
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = None  # Default color
    
    # Add horizontal line after title
    doc.add_paragraph("_" * 80)
    
    # BODY
    if body:
        body_paragraphs = body.split('\n\n')
        for para_text in body_paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.style.font.size = Pt(11)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # FOOTER
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Confidential Document | PNB Housing Finance Ltd."
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to bytes if no output path specified
    if output_path:
        doc.save(output_path)
        return output_path
    else:
        # Return as bytes for web response
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()

if __name__ == "__main__":
    generate_docx()